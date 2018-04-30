#!/usr/bin/python
# -*- coding: utf-8 -*-
from docx import Document

from tkinter import *
from tkinter import ttk

import tkSimpleDialog as tkSD
import tkFileDialog
import tkMessageBox as tkMB

import legal

import os


class mainapp:

    def __init__(self, master, *args, **kwargs):
        self.master = master
        master.title("Affidavit Generator")

        # top menubar
        self.menubar = Menu(root)
        self.filemenu = Menu(self.menubar, tearoff=0)
        self.filemenu.add_command(label="About", command=self.about)
        self.filemenu.add_separator()
        self.filemenu.add_command(label="Exit", command=root.quit)
        self.menubar.add_cascade(label="File", menu=self.filemenu)
        root.config(menu=self.menubar)

        self.title = Label(root, text="What you'll need:")
        self.title.grid(row=2, column=0, padx=4, pady=4, sticky="NW")

        # frame for holding required information
        self.info_frame = Frame(root, bd=0)
        self.info_frame.grid(row=2, column=1, padx=4, pady=4, sticky="W")
        self.info_frame['borderwidth'] = 2
        self.info_frame['relief'] = 'sunken'

        # text box to insert required information
        self.req_info = StringVar()
        self.req_info = Text(self.info_frame, width=40, height=4, wrap='word')
        self.req_info.grid(row=2, column=1)
        self.req_info.config(state='disabled')

        # scrollbar for required information text box
        self.req_scroll = Scrollbar(self.info_frame,
                                    command=self.req_info.yview)
        self.req_scroll.grid(row=2, column=1, sticky='nsew')
        self.req_info['yscrollcommand'] = self.req_scroll.set

        # combobox for selecting affidavit
        self.label = Label(root, text="Affidavit Type",
                           padx=4, pady=4)
        self.label.grid(row=1, column=0, sticky="NW")

        self.value = StringVar()
        self.afftype = ttk.Combobox(self.master, textvariable=self.value,
                                    state='readonly')
        self.afftype['values'] = aff_info_required.keys()
        self.afftype.grid(row=1, column=1)
        self.afftype.bind('<<ComboboxSelected>>', self.displayinfo)
        self.afftype.current(0)
        self.displayinfo(self)

        # button for running python-docx script
        self.make = Button(root, text="Generate", command=self.make_doc)
        self.make.grid(row=7, column=1, padx=10, pady=10)

        # directory select button
        self.save_city = Button(root, text='Browse',
                                command=self.set_directory)
        self.save_city.grid(row=5, column=5, padx=10, pady=10)

        self.save_label = Label(root, text="Save location:")
        self.save_label.grid(row=5, column=0, padx=4, pady=4, sticky="W")

        self.start_dir = StringVar()
        self.start_dir.set(os.getcwd())

        self.dir_label = Text(root, width=40, height=1, state='normal',
                              wrap='word')
        self.dir_label.insert('1.0', self.start_dir.get())
        self.dir_label.config(state='disabled')
        self.dir_label.grid(row=5, column=1)

    def client_exit(self):
        exit()

    def set_directory(self):
        tmp_dir = tkFileDialog.askdirectory()
        self.dir_label.config(state='normal')
        self.dir_label.delete('1.0', END)
        self.dir_label.insert('1.0', tmp_dir)
        self.dir_label.config(state='disabled')

    def get_directory(self):
        current_dir = self.dir_label.get('1.0', 'end-1c')
        return current_dir

    def about(self):
        tkMB.showinfo("About",
                      "made by Shaun Anderson with python-docx")

    def displayinfo(self, event):
        self.req_info.config(state='normal')
        self.req_info.delete('1.0', END)

        print "User selection is %s" % self.afftype.get()

        length = (len(aff_info_required[self.afftype.get()]))
        y = 1
        for i in aff_info_required[self.afftype.get()]:
            if y != length:
                self.req_info.insert('1.0', i + '\n')
            else:
                self.req_info.insert(END, i)
            y += 1

        self.req_info.config(state='disabled')

    def get_info(self):

            self.doc_type = self.afftype.get()

            self.stName = tkSD.askstring("Student Name",
                                         "Please enter the student's name")

            if (self.doc_type == 'Separated' or
               self.doc_type == 'Common Law'):

                self.spName = tkSD.askstring("Spouse Name",
                                             "Please enter the spouse's name")
            else:
                self.spName = ''

            if self.doc_type == 'Separated' or 'Sole Support':
                self.streetname = tkSD.askstring("Civic Addres",
                                                 legal.q_street)
            else:
                self.streetname = ''

            self.city = tkSD.askstring("city",
                                       "Please enter the city and province")

            if self.doc_type == 'Common Law':
                self.livingsince = tkSD.askstring("Date Living Since",
                                                  legal.q_date_living)
                self.separated = ''

            elif self.doc_type == 'Separated':
                    self.livingsince = ''
                    self.separated = tkSD.askstring("Date Separated",
                                                    ("Please enter the date "
                                                     "the couple separated"))
            elif self.doc_type == 'Sole Support':
                self.livingsince = ''
                self.separated = ''

            self.children = tkMB.askquestion("Children",
                                             "Do they have dependents?")

            self.chNames = []
            self.chBdays = []

            if self.children == 'yes':
                self.number_of_children = tkSD.askinteger("Number",
                                                          "How many kids?")
                for i in range(0, self.number_of_children):
                    self.chNames.append(tkSD.askstring("Name", legal.q_chn
                                                       % (i + 1)))

                    self.chBdays.append(tkSD.askstring("BDAY", (legal.q_ch
                                                       % (i + 1))))
            else:
                self.number_of_children = 0

            if ((self.doc_type == 'Separated' or
               self.doc_type == 'Sole Support') and self.children == 'yes'):

                self.open_cst_window()

                if self.doc_type == 'Separated':
                    self.st_custody = self.app.selection
                    self.sole_status = ''

                elif self.doc_type == 'Sole Support':
                    self.sole_status = self.app.selection
                    self.st_custody = ''

            else:
                self.st_custody = ''
                self.sole_status = ''

            return {'Student Name':
                    self.stName,
                    'Spouse Name':
                    self.spName,
                    'city':
                    self.city,
                    'Living Since':
                    self.livingsince,
                    'Children Names':
                    self.chNames,
                    'Children Birthdays':
                    self.chBdays,
                    'Children':
                    self.number_of_children,
                    'Custody':
                    self.st_custody,
                    'Sole Status':
                    self.sole_status,
                    'Separated Date':
                    self.separated,
                    'Street':
                    self.streetname
                    }

    def open_cst_window(self):

        self.cst_window = Toplevel(self.master)
        self.app = Custody(self.cst_window, self.doc_type)
        self.cst_window.geometry('325x120+412+342')
        self.master.wait_window(self.cst_window)

    def make_doc(self):
        # grab user selected directory and affidavit type
        self.get_directory()
        self.answers = self.get_info()
        print self.answers

        document = Document('affidavit-template.docx')

        # save the file
        filename = str("Affidavit-%s.docx" % self.answers["Student Name"])
        filedirectory = self.get_directory()
        filepath = os.path.join(filedirectory, filename)
        document.save(filepath)

        document._body.clear_content()

        # affidavit header with scilicet
        title = document.add_heading('AFFIDAVIT')
        title.style = document.styles['Heading 1']
        run = title.add_run()
        run.add_break()

        country = document.add_table(rows=4, cols=2)
        country.style = 'none'

        canada = country.cell(0, 0)
        canada.text = 'Canada'

        rightBracket = country.cell(0, 1)
        rightBracket.text = ')'

        province = country.cell(1, 0)
        province.text = 'Province of Ontario'

        provinceRB = country.cell(1, 1)
        provinceRB.text = ') S. S.'

        finalRB = country.cell(2, 1)
        finalRB.text = ')'

        # solemny affirm and declare
        if self.afftype.get() == 'Separated':

            p = document.add_paragraph(legal.sole_declare %
                                       (self.answers["Student Name"],
                                        self.answers["city"]))
            run = p.add_run()
            run.add_break()

        elif self.afftype.get() == 'Sole Support':

            p = document.add_paragraph(legal.sole_declare %
                                       (self.answers["Student Name"],
                                        self.answers["city"]))
            run = p.add_run()
            run.add_break()

        elif self.afftype.get() == 'Common Law':
            p = document.add_paragraph(legal.declare %
                                       (self.answers["Student Name"],
                                        self.answers["Spouse Name"],
                                        self.answers["city"]))
            run = p.add_run()
            run.add_break()

        # for the purpose of supporting an OSAP application
        p = document.add_paragraph(legal.aff_purpose)
        p.style = document.styles['ListNumber']
        run = p.add_run()
        run.add_break()

        if self.afftype.get() == 'Common Law':

            p = document.add_paragraph(legal.living %
                                       self.answers["Living Since"])
            p.style = document.styles['ListNumber']
            run = p.add_run()
            run.add_break()

        elif self.afftype.get() == 'Separated':
            p = document.add_paragraph(legal.separated)
            p.style = document.styles['ListNumber']
            run = p.add_run()
            run.add_break()

            p = document.add_paragraph("The name of my spouse is %s."
                                       % self.answers["Spouse Name"])
            p.style = document.styles['ListNumber']
            run = p.add_run()
            run.add_break()

            p = document.add_paragraph(legal.separated_date
                                       % self.answers["Separated Date"])
            p.style = document.styles['ListNumber']
            run = p.add_run()
            run.add_break()

        elif self.afftype.get() == 'Sole Support':

            if self.answers['Sole Status'] == 'Widowed':
                p = document.add_paragraph(legal.widowed)
                run = p.add_run()
                run.add_break()

            elif self.answers['Sole Status'] == 'Never Married':
                p = document.add_paragraph(legal.never_m)
                run = p.add_run()
                run.add_break()

            p.style = document.styles['ListNumber']
            run = p.add_run()

        if self.answers["Children"] > 0:
            if self.afftype.get() == 'Common Law':
                p = document.add_paragraph(legal.cl_cust
                                           % self.answers['Children'])
            elif self.afftype.get() == 'Separated':
                p = document.add_paragraph(legal.sp_cust
                                           % self.answers['Children'])
            elif self.afftype.get() == 'Sole Support':
                p = document.add_paragraph(legal.sole_pnt
                                           % self.answers['Children'])

            p.style = document.styles['ListNumber']
            run = p.add_run()
            run.add_break()
            run.add_break()

            children_names = self.answers['Children Names']
            children_bdays = self.answers['Children Birthdays']

            for i in range(0, self.answers["Children"]):
                    run = p.add_run("%s, born %s" %
                                    (children_names[i], children_bdays[i]))
                    run.add_break()

        if self.afftype.get() == 'Separated':

            if self.answers['Custody'] == 'Sole Support':

                p = document.add_paragraph(legal.solecust)
                p.style = document.styles['ListNumber']
                run = p.add_run()
                run.add_break()
                run.add_break()

                run = p.add_run(self.answers['Street'])
                run.add_break()
                run = p.add_run(self.answers['city'])
                run.add_break()

            elif self.answers['Custody'] == 'Shared Custody':
                p = document.add_paragraph(legal.sharecust)
                p.style = document.styles['ListNumber']
                run = p.add_run()
                run.add_break()
                run.add_break()

                run = p.add_run('[DETAILS]')
                run.add_break()

        if self.afftype.get() == 'Sole Support':

            p = document.add_paragraph(legal.solecust)
            p.style = document.styles['ListNumber']
            run = p.add_run()
            run.add_break()
            run.add_break()

            run = p.add_run(self.answers['Street'])
            run.add_break()
            run = p.add_run(self.answers['city'])
            run.add_break()

        if self.afftype.get() == 'Common Law':

            p = document.add_paragraph(legal.marriage_law)
            p.style = document.styles['ListNumber']

            run = p.add_run()
            run.add_break()
            run.add_break()
            run.add_break()

        elif (self.afftype.get() == 'Separated' or
              self.afftype.get() == 'Sole Support'):

            p = document.add_paragraph(legal.s_oath)
            p.style = document.styles['ListNumber']

            run = p.add_run()
            run.add_break()
            run.add_break()

        # table to hold signatures
        table = document.add_table(rows=6, cols=2)
        table.style = 'Hello'

        sworn = table.cell(0, 0)
        sworn.text = legal.affirmed

        if self.afftype.get() == 'Common Law':
            firstSig = table.cell(0, 1)
            firstUnderline = table.cell(0, 1)
            firstUnderline.text = legal.underline

        elif (self.afftype.get() == 'Separated' or
              self.afftype.get() == 'Sole Support'):

            firstSig = table.cell(4, 1)
            firstUnderline = table.cell(4, 1)
            firstUnderline.text = legal.underline

        firstSig.add_paragraph(self.answers["Student Name"])

        commUnderline = table.cell(5, 0)
        commUnderline.add_paragraph(legal.underline)

        commissioner = table.cell(5, 0)
        commissioner.add_paragraph("")

        if self.afftype.get() == 'Common Law':

            secondUnderline = table.cell(5, 1)
            secondUnderline.add_paragraph(legal.underline)

            secondSig = table.cell(5, 1)
            secondSig.add_paragraph(self.answers["Spouse Name"])

        # save the file
        filename = str("Affidavit-%s.docx" % self.answers["Student Name"])
        filedirectory = self.get_directory()
        filepath = os.path.join(filedirectory, filename)

        document.save(filepath)


class Custody:

    def __init__(self, master, doc_type):

        if doc_type == 'Separated':
            self.master = master
            master.title("Custody Details")
            self.msg = Label(master,
                             text=("Is the custody shared "
                                   "or does the student have sole support?"))
            self.msg.grid(row=0, column=0, padx=5, pady=5)

            # create listbox of children names entered
            self.cust = Listbox(master, selectmode=SINGLE)
            self.cust.config(height=2)
            self.cust.grid(row=1, column=0, padx=5, pady=5)

            self.cust.insert(END, "Shared Custody")
            self.cust.insert(END, "Sole Support")

            self.btn = Button(master, text="Okay",
                              command=self.select)
            self.btn.grid(row=2, column=0)

        elif doc_type == 'Sole Support':
            self.master = master
            master.title("Sole Support Status")
            self.msg = Label(master,
                             text=("Please select the appropriate status"))
            self.msg.grid(row=0, column=0, padx=5, pady=5)

            # create listbox of children names entered
            self.st_sole = Listbox(master, selectmode=SINGLE)
            self.st_sole.config(height=2)
            self.st_sole.grid(row=1, column=0, padx=5, pady=5)

            self.st_sole.insert(END, "Never Married")
            self.st_sole.insert(END, "Widowed")

            self.btn = Button(master, text="Okay",
                              command=self.select)
            self.btn.grid(row=2, column=0)

    def select(self):

        if app.doc_type == 'Separated':
            self.selection = self.cust.get(self.cust.curselection())

        elif app.doc_type == 'Sole Support':
            self.selection = self.st_sole.get(self.st_sole.curselection())

        self.master.destroy()


if __name__ == '__main__':

    root = Tk()

    aff_info_required = {'Common Law':
                         ['Student Name', 'Spouse Name',
                          'Number of Children', 'Names of Children',
                          'Birthdates of Children', 'City, Province',
                          'Date living since'],
                         'Separated':
                         ['Student Name', 'Spouse Name',
                          'Number of Children', 'Names of Children',
                          'Birthdates of Children', 'Address',
                          'Custody Shared or Sole Support'],
                         'Sole Support':
                         ['Student Name', 'Never Married or Widowed',
                          'Number of Children', "Birthdates of Children",
                          'Address'],
                         }

    w = 550
    h = 200

    # get screen width and height
    ws = root.winfo_screenwidth()
    hs = root.winfo_screenheight()

    # calculate x and y coordinates for the TK root window
    x = (ws / 2) - (w - 2)
    y = (hs / 2) - (h - 2)

    app = mainapp(root)
    root.geometry('%dx%d+%d+%d' % (w, h, x, y))
    root.resizable(False, False)
    root.mainloop()

from docx import Document
from docx.shared import Pt
from Tkinter import *
import time
import datetime

document = Document('affidavit-template.docx')
document.save('affidavit-new.docx')
document._body.clear_content()

fields = 'Student Name', 'Spouse Name', 'Location', 'Date Living Since'

def makeform(root, fields):
   entries = []
   for field in fields:
      row = Frame(root)
      lab = Label(row, width=15, text=field, anchor='w')
      ent = Entry(row)
      row.pack(side=TOP, fill=X, padx=5, pady=5)
      lab.pack(side=LEFT)
      ent.pack(side=RIGHT, expand=YES, fill=X)
      entries.append((field, ent))
   return entries


if __name__ == '__main__':
   root = Tk()
   ents = makeform(root, fields)
   root.bind('<Return>', (lambda event, e=ents: fetch(e)))
   b1 = Button(root, text='Generate')
   b1.pack(side=LEFT, padx=5, pady=5)
   b2 = Button(root, text='Quit', command=root.quit)
   b2.pack(side=LEFT, padx=5, pady=5)
   root.mainloop()



i = time.strftime("%c")
now = datetime.datetime.now()

underline = "_______________________________________"

for style in document.styles:
    print style.name

#print "What is the student's first and last name?"
#studentName = raw_input()

#print "What is the spouse's first and last name?"
#spouseName = raw_input()

#print "What is the city and province?"
#location = raw_input()

#print "What date did the couple start living together?"
#livingsince = raw_input()

title = document.add_heading('AFFIDAVIT OF MARRIAGE')
title.style = document.styles['Heading 1']
run = title.add_run()
run.add_break()

p = document.add_paragraph(
    "We, %s and %s, both of %s, SOLEMNLY AFFIRM AND DECLARE THAT:" % (
    studentName, spouseName, location))

run = p.add_run()
run.add_break()

p = document.add_paragraph(
        "We are living together in a conjugal relationship, and have done so continuously since %s""" % livingsince)
p.style = document.styles['ListNumber']
run = p.add_run()
run.add_break()

p = document.add_paragraph(
        "The information contained in this affidavit is provided for the sole purpose of supporting an application for financial aid under the Ontario Student Assistance Program (OSAP).")
p.style = document.styles['ListNumber']
run = p.add_run()
run.add_break()

p = document.add_paragraph(
        "We understand that the acceptance of this affidavit by the Carleton University Awards Office and the Ontario Ministry of Advanced Education and Skills Development in no way constitutes a legal determination that our common law marriage is valid under applicable law.")
p.style = document.styles['ListNumber']
run = p.add_run()
run.add_break()
run.add_break()
run.add_break()

table = document.add_table(rows = 6, cols = 2)
table.style = 'Hello'

sworn = table.cell(0, 0)
sworn.add_paragraph("SWORN/AFFIRMED BEFORE ME, at Carleton University, Ottawa, Ontario this %s day of %s, %s" % (time.strftime("%d"), time.strftime("%B"), now.year))

firstUnderline = table.cell(0, 1)
firstUnderline.add_paragraph(underline)

firstSig = table.cell(0, 1)
firstSig.add_paragraph(studentName)

commUnderline = table.cell(5, 0)
commUnderline.add_paragraph(underline)

commissioner = table.cell(5, 0)
commissioner.add_paragraph("Mark Alexander Robinson")

secondUnderline = table.cell(5, 1)
secondUnderline.add_paragraph(underline)

secondSig = table.cell(5, 1)
secondSig.add_paragraph(spouseName)

document.save('affidavit-new.docx')

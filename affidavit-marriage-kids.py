from docx import Document
from docx.shared import Pt
import legal
import Tkinter
import tkSimpleDialog


document = Document('affidavit-template.docx')
document.save('affidavit-new.docx')
document._body.clear_content()

root = Tkinter.Tk()
root.geometry("1000x300")
root.withdraw()

childrenNames = []
childrenBdays = []

studentName = tkSimpleDialog.askstring("Student Name", "Please enter the student's name")
spouseName =  tkSimpleDialog.askstring("Spouse Name", "Please enter the spouse's name")
location = tkSimpleDialog.askstring("Location", "Please enter the city and province")
livingsince = tkSimpleDialog.askstring("Date Living Since", "Please enter the date the couple began living together")
children = tkSimpleDialog.askinteger("Number of Chidlren", "Please enter the number of children in number format")

#appends children name and birth dates to separate lists to be later inserted into table
for i in range(0, children):
    childrenNames.append(tkSimpleDialog.askstring(
                    "Child Name", "Please enter the name of child %d" % (i+1)))
    childrenBdays.append(tkSimpleDialog.askstring("Child Birth Date", "Please enter the birth date of child %d in MONTH DATE, YEAR format" % (i+1)))

title = document.add_heading('AFFIDAVIT OF MARITAL STATUS')
title.style = document.styles['Heading 1']
run = title.add_run()
run.add_break()

country = document.add_table(rows = 4, cols = 2)
country.style = 'none'

canada = country.cell(0,0)
canada.text = 'Canada'

rightBracket = country.cell(0,1)
rightBracket.text = ')'

province = country.cell(1,0)
province.text = 'Province of Ontario'

provinceRB = country.cell (1,1)
provinceRB.text = ') S. S.'

finalRB = country.cell (2,1)
finalRB.text = ')'

p = document.add_paragraph(
"We, %s and %s, both of %s, SOLEMNLY AFFIRM AND DECLARE THAT:" % (
studentName, spouseName, location))

run = p.add_run()
run.add_break()

p = document.add_paragraph(
        "We are living together in a conjugal relationship, and have done so continuously since %s""" % livingsince)
p.style = document.styles['ListNumber']
run = p.add_run()
run.add_break()

p = document.add_paragraph(
        "We are the custodial and natural (or adoptive) parents of %d child(ren), namely:" % children)
p.style = document.styles['ListNumber']
run = p.add_run()
run.add_break()
run.add_break()

for i in range(0, children):
        run = p.add_run("%s, born %s" % (childrenNames[i], childrenBdays[i]))
        run.add_break()
        run.add_break()

p = document.add_paragraph(
        "The information contained in this affidavit is provided for the sole purpose of supporting an application for financial aid under the Ontario Student Assistance Program (OSAP).")
p.style = document.styles['ListNumber']
run = p.add_run()
run.add_break()

p = document.add_paragraph(
        "We understand that the acceptance of this affidavit by the Carleton University Awards Office and the Ontario Ministry of Advanced Education and Skills Development in no way constitutes a legal determination that our common law marriage is valid under applicable law.")
p.style = document.styles['ListNumber']
run = p.add_run()
run.add_break()
run.add_break()
run.add_break()

table = document.add_table(rows = 6, cols = 2)
table.style = 'Hello'

sworn = table.cell(0, 0)
sworn.text = legal.affirmed

firstUnderline = table.cell(0, 1)
firstUnderline.text = legal.underline

firstSig = table.cell(0, 1)
firstSig.add_paragraph(studentName)

commUnderline = table.cell(5, 0)
commUnderline.add_paragraph(legal.underline)

commissioner = table.cell(5, 0)
commissioner.add_paragraph("Mark Alexander Robinson")

secondUnderline = table.cell(5, 1)
secondUnderline.add_paragraph(legal.underline)

secondSig = table.cell(5, 1)
secondSig.add_paragraph(spouseName)

document.save('affidavit-%s.docx' % studentName)
